from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_DIR = DOCS_DIR / "sna_project_outputs"
TABLE_DIR = OUTPUT_DIR / "tables"
FIGURE_DIR = OUTPUT_DIR / "figures"
REPORT_PATH = DOCS_DIR / "sna_project_report.docx"


ACCENT = RGBColor(28, 63, 95)
BODY = RGBColor(34, 40, 49)
MUTED = RGBColor(90, 98, 104)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = BODY

    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = ACCENT
        styles[style_name].font.size = Pt(size)


def add_paragraph(document: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = BODY


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def add_image(document: Document, path: Path, caption: str, width_inches: float = 6.8) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Inches(width_inches))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(document, caption)


def dataframe_to_table(document: Document, dataframe: pd.DataFrame, title: str | None = None) -> None:
    if title:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(title)
        run.bold = True
        run.font.color.rgb = ACCENT

    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for index, column in enumerate(dataframe.columns):
        header_cells[index].text = str(column)
        for paragraph in header_cells[index].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        header_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(header_cells[index], "1C3F5F")

    for row in dataframe.itertuples(index=False):
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = str(value)
            row_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in row_cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph()


def load_metrics() -> dict:
    return json.loads((TABLE_DIR / "network_metrics.json").read_text(encoding="utf-8"))


def build_report() -> Path:
    metrics = load_metrics()
    centrality = pd.read_csv(TABLE_DIR / "centrality_top5.csv")
    community_summary = pd.read_csv(TABLE_DIR / "community_summary.csv")
    robustness = pd.read_csv(TABLE_DIR / "robustness_summary.csv")

    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Sosyal Ag Analizi Tabanli Urun-Influencer Esleme Sistemi")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Creator similarity graph uzerinden merkezilik, topluluk ve dayanıklılık analizi ile urun odakli influencer onerisi"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = MUTED

    add_paragraph(
        document,
        "Bu rapor, sosyal ag analizi dersinin ag modelleme, temel olcutler, merkezilik, topluluk analizi, dayaniklilik ve yorumlama isterlerini, "
        "urun sorgusuna gore influencer onerisi yapan uygulamali bir creator benzerlik agi uzerinden karsilamaktadir.",
    )

    document.add_heading("1. Konu ve Arastirma Sorusu", level=1)
    add_paragraph(
        document,
        "Projede analiz edilen ag, influencer hesaplari arasindaki kategori ve hashtag benzerliginden uretilen agirlikli creator similarity graph yapisidir. "
        "Dugumler tekil creator hesaplarini, kenarlar ise iki creator arasindaki icerik yakinligini temsil etmektedir. Ag yonlusuz ve agirliklidir; "
        "agirlik degeri kategori benzerligi ile hashtag benzerliginin birlesiminden hesaplanmistir.",
    )
    add_paragraph(
        document,
        "Temel arastirma sorusu sudur: Bir urun sorgusu girildiginde, alan otoritesi, etkilesim gucu ve gerekirse kopru rolu dikkate alinarak hangi influencer "
        "daha mantikli bir aday olarak onerilmelidir?",
    )
    add_image(
        document,
        DOCS_DIR / "excalidraw" / "exports" / "03_gercek_graph_modeli.png",
        "Sekil 1. Projede kullanilan gercek graph modeli ve dugum/kenar mantigi.",
        6.6,
    )

    document.add_heading("2. Veri Seti ve Etik Uygunluk", level=1)
    add_paragraph(
        document,
        "Veri kaynagi Instagram profil ve hashtag taramalarindan turetilen yerel veri setidir. Analizde kullanilan ana dosya "
        "clean_creator_dataset.csv olup, ham tarama dosyalari ve islenmis veri katmani proje klasorunde saklanmistir. Calisma yalnizca acik profil ve icerik "
        "sinyallerine dayali egitsel bir prototip niteligindedir; ozel mesaj, gizli profil veya hassas kisisel veri kullanilmamistir.",
    )

    dataset_table = pd.DataFrame(
        [
            ["Veri kaynagi", "Instagram profil/hashtag taramasi, yerel islenmis veri"],
            ["Dugum turu", "Creator / influencer hesabi"],
            ["Kenar turu", "Icerik benzerligi (kategori + hashtag overlap)"],
            ["Dugum sayisi", metrics["node_count"]],
            ["Kenar sayisi", metrics["edge_count"]],
            ["Veri formati", "CSV, JSON, GraphML, PNG, DOCX, IPYNB, Python"],
        ],
        columns=["Baslik", "Aciklama"],
    )
    dataframe_to_table(document, dataset_table)

    document.add_heading("3. Ag Modelleme", level=1)
    add_paragraph(
        document,
        "Ag G = (V, E) seklinde tanimlanmistir. V creator hesaplarini, E ise iki creator arasindaki benzerlik kenarlarini temsil eder. "
        "Projede dugum listesi, kenar listesi, komsuluk matrisi, komsuluk listesi ve GraphML ciktisi uretilmistir. Bu sayede ag hem Python/NetworkX tarafinda "
        "hem de Gephi gibi araclarda tekrar incelenebilir durumdadir.",
    )
    add_bullet(document, "Dugum listesi: docs/sna_project_outputs/tables/node_list.csv")
    add_bullet(document, "Kenar listesi: docs/sna_project_outputs/tables/edge_list.csv")
    add_bullet(document, "Komsuluk matrisi: docs/sna_project_outputs/tables/adjacency_matrix.csv")
    add_bullet(document, "GraphML cikisi: docs/sna_project_outputs/creator_similarity_graph.graphml")
    add_image(
        document,
        DOCS_DIR / "excalidraw" / "exports" / "02_graf_icin_veri_hazirlama.png",
        "Sekil 2. Ham veri katmanindan ag modeline gecis akisi.",
        6.6,
    )

    document.add_heading("4. Temel Ag Olcutleri", level=1)
    metrics_table = pd.DataFrame(
        [
            ["Dugum sayisi", metrics["node_count"]],
            ["Kenar sayisi", metrics["edge_count"]],
            ["Bagli bilesen sayisi", metrics["connected_components"]],
            ["En buyuk bilesen", f'{metrics["largest_component_size"]} dugum ({metrics["largest_component_share"] * 100:.2f}%)'],
            ["Yogunluk", metrics["density"]],
            ["Ortalama derece", metrics["average_degree"]],
            ["Ortalama agirlikli derece", metrics["average_weighted_degree"]],
            ["Ortalama kumeleme katsayisi", metrics["average_clustering"]],
            ["Cap (LCC)", metrics["diameter_lcc"]],
            ["Ortalama en kisa yol (LCC)", metrics["average_shortest_path_lcc"]],
        ],
        columns=["Olcut", "Deger"],
    )
    dataframe_to_table(document, metrics_table)
    add_paragraph(
        document,
        "Yogunluk degeri 0.0889 oldugu icin ag ne tamamen seyrek ne de asiri yogundur; orta yogunlukta bir benzerlik omurgasi sunmaktadir. "
        "Cap degerinin 6 olmasi, en buyuk bagli bilesende creatorlar arasindaki yapisal uzakligin gorece sinirli oldugunu, yani kompakt bir cekirdek yapı "
        "bulundugunu gostermektedir. Kumeleme katsayisinin yuksek olmasi ise benzer ilgi alanlarina sahip creatorlarin kendi iclerinde grup olusturma egilimini "
        "destekler.",
    )
    add_image(
        document,
        FIGURE_DIR / "05_network_overview_polished.png",
        "Sekil 3. Genel ag yapisi. Renkler baskin kategori ailelerini, dugum buyuklugu dereceyi temsil eder.",
        6.7,
    )
    add_image(
        document,
        FIGURE_DIR / "08_degree_distribution_polished.png",
        "Sekil 4. Derece dagilimi. Sag kuyruk yapisi, hub niteligindeki dugumlerin varligini gosterir.",
        6.3,
    )

    document.add_heading("5. Merkezilik Analizi", level=1)
    add_paragraph(
        document,
        "Merkezilik yorumlari en buyuk bagli bilesen uzerinden yapilmistir. Degree centrality en baglantili dugumleri, betweenness centrality farkli "
        "topluluklari birlestiren kopru hesaplari, closeness centrality ag icinde hizli erisebilen dugumleri, eigenvector ve PageRank ise etkili ve "
        "otorite benzeri dugumleri one cikarir.",
    )
    for measure in [
        "Degree Centrality",
        "Betweenness Centrality",
        "Closeness Centrality",
        "Eigenvector Centrality",
        "PageRank",
    ]:
        subset = centrality[centrality["measure"] == measure][["rank", "username", "score", "categories"]]
        dataframe_to_table(document, subset, title=f"{measure} icin ilk 5 dugum")
    add_paragraph(
        document,
        "Sonuclar birlikte okundugunda annevebebek2025 ve benzeri hesaplar yuksek baglanti ve otorite rollerinde one cikarken, studyconkris gibi hesaplar "
        "farkli icerik topluluklari arasinda gecis saglayan kopru rolu ustlenmektedir. Bu yorum, bizim agimiz takip iliskisi degil benzerlik iliskisi uzerine "
        "kuruldugu icin 'kopru' kavraminin da topluluklar arasi icerik baglantisi olarak okunmasini gerektirir.",
    )
    add_image(
        document,
        FIGURE_DIR / "12_centrality_focus_subgraph.png",
        "Sekil 5. Betweenness odakli alt ag. Buyuk dugumler yapisal kopru rolu tasiyan hesaplari vurgular.",
        6.8,
    )
    add_image(
        document,
        FIGURE_DIR / "10_top_centrality_comparison.png",
        "Sekil 6. Farkli merkezilik olcutlerinde one cikan dugumlerin karsilastirmasi.",
        6.6,
    )

    document.add_heading("6. Topluluk Analizi", level=1)
    add_paragraph(
        document,
        "Topluluk analizi Louvain algoritmasi ile gerceklestirilmistir. Ag 8 topluluga ayrilmistir ve modularity degeri 0.630465'tir. Bu yuksek deger, "
        "topluluklarin rastgele degil anlamli kategori yogunlasmalari etrafinda toplandigini gostermektedir.",
    )
    dataframe_to_table(
        document,
        community_summary[["community_id", "size", "share_of_lcc", "top_categories"]],
        title="Topluluk ozet tablosu",
    )
    add_paragraph(
        document,
        "En buyuk topluluk otomotiv, guzellik_kozmetik ve oyuncak_hobi ekseninde genis bir yapi sunarken; ikinci ve ucuncu topluluklar sirasiyla yemek_mutfak "
        "ve egitim_kariyer temalarinda yogunlasmaktadir. Topluluklar arasi baglari guclendiren dugumler, ozellikle cok kategorili urun senaryolarinda bridge "
        "modunun neden onemli oldugunu gostermektedir.",
    )
    add_image(
        document,
        FIGURE_DIR / "07_louvain_communities_polished.png",
        "Sekil 7. Louvain topluluklari. Renkler farkli topluluklari, secili etiketler yapisal olarak onemli dugumleri gosterir.",
        6.8,
    )
    add_image(
        document,
        FIGURE_DIR / "13_community_meta_graph.png",
        "Sekil 8. Topluluklar arasi meta-graph. Dugum buyuklugu topluluk buyuklugunu, kenarlar topluluklar arasi bag gucunu temsil eder.",
        6.3,
    )
    add_image(
        document,
        FIGURE_DIR / "09_community_size_chart.png",
        "Sekil 9. Topluluk buyukluklerinin dagilimi.",
        6.0,
    )

    document.add_heading("7. Gorsellestirme ve Ornek Ciktilar", level=1)
    add_paragraph(
        document,
        "Projede agi farkli acilardan okunabilir kilmak icin hem tum-ag gorselleri hem de odakli analitik gorseller uretildi. Genel bakis, merkezilik odagi, "
        "topluluk yapisi, derece dagilimi, meta-graph ve dayanıklılık grafikleri birlikte kullanildiginda agin hem yerel hem de makro davranisi daha acik "
        "sekilde yorumlanabilmektedir.",
    )
    add_image(
        document,
        FIGURE_DIR / "11_backbone_overview.png",
        "Sekil 10. Yalnizca guclu benzerlik kenarlarini gosteren omurga grafigi.",
        6.6,
    )

    document.add_heading("8. Dayaniklilik ve Ileri Analiz", level=1)
    dataframe_to_table(document, robustness, title="Dayaniklilik senaryolari")
    add_paragraph(
        document,
        "Tek bir yuksek degree dugumunun cikmasi agi dramatik bicimde parcalamamistir. Ancak yuksek betweenness dugumlerinin hedefli bicimde kaldirilmasi, "
        "diameter degerini artirmis ve erisim verimliligini zayiflatmistir. Bu bulgu, graph tabanli oneride kopru dugumlerin neden stratejik onem tasidigini "
        "dogrudan gosterir.",
    )
    add_image(
        document,
        FIGURE_DIR / "14_robustness_comparison.png",
        "Sekil 11. Rastgele ve hedefli dugum cikarma senaryolarinin dayanıklılık uzerindeki etkisi.",
        6.8,
    )

    document.add_heading("9. Bulgularin Yorumlanmasi: Urun-Influencer Onerisi", level=1)
    add_paragraph(
        document,
        "Projenin uygulamali katmaninda urun sorgusu once urun sinyallerine ayrilmakta, daha sonra graph uzerindeki authority_score, bridge_score, "
        "engagement_score ve follower_score ile birlestirilmektedir. Tek kategorili senaryolarda authority modu, cok kategorili senaryolarda ise bridge modu "
        "daha baskin rol oynar. Boylece sistem yalnizca kelime eslestirmesi yapmaz; graph yapisinin sagladigi yapisal bilgiyle secimi guclendirir.",
    )
    add_image(
        document,
        DOCS_DIR / "excalidraw" / "exports" / "04_analiz_ve_kullanim_mantigi.png",
        "Sekil 12. Urun sinyalleri ile graph olcutlerinin oneride birlesme mantigi.",
        6.8,
    )
    add_image(
        document,
        FIGURE_DIR / "15_recommendation_examples.png",
        "Sekil 13. Ornek urun sorgulari icin ilk 5 influencer onerisi ve skor dagilimlari.",
        6.8,
    )

    document.add_heading("10. Kod, Notebook ve Teslim Duzeni", level=1)
    add_paragraph(
        document,
        "Teslim paketi; Python kodlari, Jupyter Notebook, veri seti, ag gorselleri, GraphML cikisi ve rapor dosyasini birlikte icermektedir. "
        "sna_project.py analiz omurgasini, recommender.py oneriyi, taxonomy.py urun niyeti cikarmayi, app.py ise temel arayuzu yonetmektedir. "
        "Notebook dosyasi tekrar edilebilir analiz akisini sunarken, gorseller rapor ve sunumda kullanilmak uzere ayri klasorde saklanmistir.",
    )
    add_bullet(document, "Python kodlari: sna_project.py, recommender.py, taxonomy.py, app.py")
    add_bullet(document, "Notebook: docs/sna_creator_network.ipynb")
    add_bullet(document, "Veri seti: data/processed/clean_creator_dataset.csv")
    add_bullet(document, "Gorseller ve tablolar: docs/sna_project_outputs/")

    document.add_heading("11. Sonuc", level=1)
    add_paragraph(
        document,
        "Calisma sonucunda creator similarity graph yapisinin urun odakli influencer onerisi icin guclu bir karar omurgasi sundugu gorulmustur. "
        "Ag orta yogunlukta, yuksek kumeleme egilimli ve anlamli topluluklara ayrilan bir yapidadir. Merkezi dugumler ile kopru dugumler ayni degildir; "
        "bu nedenle influencer secimi tek bir olcege baglanmamis, senaryoya gore authority ve bridge dengesi kurulmustur.",
    )
    add_paragraph(
        document,
        "Temel sinirlilik, agin gercek takip iliskisi yerine benzerlik iliskisi uzerine kurulmus olmasidir. Buna ragmen merkezilik, topluluk ve dayaniklilik "
        "bulgulari hem ders isterlerini karsilamakta hem de gercek bir urun sorgusuna uygulanabilir bir onerme sistemi ortaya koymaktadir.",
    )

    document.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    output = build_report()
    print(output)
