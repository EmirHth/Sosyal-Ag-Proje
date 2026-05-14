from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_PATH = DOCS_DIR / "prezi_sunum_plani.docx"

ACCENT = RGBColor(28, 63, 95)
BODY = RGBColor(34, 40, 49)
MUTED = RGBColor(90, 98, 104)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = BODY

    for style_name, size in [("Heading 1", 18), ("Heading 2", 14)]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = ACCENT
        styles[style_name].font.size = Pt(size)


def add_paragraph(document: Document, text: str, *, center: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.italic = italic
    run.font.color.rgb = BODY if not italic else MUTED


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_image(document: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(10)
    caption_run.font.color.rgb = MUTED


def frame_section(
    document: Document,
    title: str,
    bullets: list[str],
    speech: str,
    image: Path | None = None,
    image_caption: str | None = None,
    image_width: float = 6.2,
) -> None:
    document.add_heading(title, level=1)
    for bullet in bullets:
        add_bullet(document, bullet)
    add_paragraph(document, f"Konusma notu: {speech}")
    if image and image_caption:
        add_image(document, image, image_caption, image_width)


def build_docx() -> Path:
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("7 Slaytlik Prezi Sunum Plani")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = "Aptos Display"
    title_run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Sosyal Ag Analizi Tabanli Urun-Influencer Esleme Sistemi")
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = MUTED

    add_paragraph(
        document,
        "Bu dosya Prezi'ye dogrudan icerik tasimak icin hazirlanmistir. Her frame icin kisa maddeler, konusma notu ve uygun gorsel secimi ayni yerde verilmistir.",
        center=False,
    )

    document.add_heading("Zoom Akisi", level=1)
    for item in [
        "1. Kapak ve problem",
        "2. Veri seti + urun sinyali / kategori cikarimi",
        "3. Ag modelleme + temel ag olcutleri",
        "4. Merkezilik analizi",
        "5. Topluluk + dayaniklilik analizi",
        "6. Urun odakli onerme mekanizmasi",
        "7. Ornek sorgular + sonuc",
    ]:
        add_bullet(document, item)

    frame_section(
        document,
        "Frame 1 - Kapak ve Problem",
        [
            "Influencer secimi sadece takipci sayisina gore yapilmamalidir.",
            "Bazi urunlerde alan otoritesi, bazilarinda ise kopru rolunde creator gerekir.",
            "Bu nedenle creatorlar arasi benzerlik agi kurulmustur.",
            "Arastirma sorusu: Bir urun sorgusu girildiginde sistem hangi influencer'i one cikarmalidir?",
        ],
        "Bu projede sosyal ag analizini sadece olcum yapmak icin degil, urune gore daha mantikli influencer secmek icin kullandik.",
        DOCS_DIR / "excalidraw" / "exports" / "05_kapak_proje_mantigi.png",
        "Kapak gorseli ve proje akisi",
        6.5,
    )

    frame_section(
        document,
        "Frame 2 - Veri Seti ve Urun Sinyali / Kategori Cikarimi",
        [
            "Veri kaynagi: Instagram profil ve hashtag taramasi",
            "Dugum: creator / influencer hesabi",
            "Kenar: kategori + hashtag benzerligi",
            "Dugum sayisi: 389 | Kenar sayisi: 6711",
            "Keyword mantigi: sorgu normalize edilir, kategori sozlugu ile eslestirilir, yardimci sinyaller puanlanir.",
            "Ornekler: kadin yuz serumu -> guzellik_kozmetik | spor matarasi -> fitness_saglik + yemek_mutfak | kamp termosu -> seyahat_gezi + yemek_mutfak",
        ],
        "Bu katman graph degildir; urunun ne istedigini anlamaya yarar. Keywordler, kullanim baglami ve urun tipi birlikte okunur; asil influencer secimi daha sonra creator graph uzerinde yapilir.",
        DOCS_DIR / "excalidraw" / "exports" / "02_graf_icin_veri_hazirlama.png",
        "Veri hazirlama ve urun sinyali mantigi",
        6.4,
    )

    frame_section(
        document,
        "Frame 3 - Ag Modelleme ve Temel Olcutler",
        [
            "G = (V, E)",
            "V: creator hesaplari",
            "E: iki creator arasindaki benzerlik kenarlari",
            "Ag tipi: yonlusuz ve agirlikli",
            "Yogunluk: 0.0889 | Ortalama derece: 34.50 | Kumeleme katsayisi: 0.6899 | Cap: 6",
        ],
        "Bu ag takip agi degil, similarity graph. Creatorlar arasindaki icerik yakinligini yapi olarak modelliyoruz.",
        DOCS_DIR / "sna_project_outputs" / "figures" / "17_network_overview_presentation.png",
        "Genel ag gorunumu ve aciklamasi",
        6.5,
    )
    add_image(
        document,
        DOCS_DIR / "sna_project_outputs" / "figures" / "18_degree_distribution_presentation.png",
        "Derece dagilimi ve hub yorumu",
        6.5,
    )

    frame_section(
        document,
        "Frame 4 - Merkezilik Analizi",
        [
            "Yontem: NetworkX merkezilik olcutleri",
            "annevebebek2025 degree ve PageRank'te one cikti; anne-bebek ekseninde cok baglantili ve otorite benzeri bir merkezdir.",
            "emine.bites eigenvector tarafinda one cikti; etkili dugumlere bagli bir creator konumundadir.",
            "itssyamaan closeness tarafinda one cikti; cekirdek ag icinde diger dugumlere hizli ulasabilmektedir.",
            "studyconkris betweenness'te cok yuksek cikar; egitim, ev_yasam ve evcil_hayvan eksenleri arasinda yapisal kopru rolundedir.",
        ],
        "Bizim projede Degree, Eigenvector ve PageRank birlikte okundugunda hangi creatorlarin alan icinde guclu otorite oldugu anlasiliyor. Betweenness ise farkli topluluklari baglayan creatorlari ortaya cikariyor. Bu fark, oneride authority ve bridge modlarini ayirmamizi sagladi.",
        DOCS_DIR / "sna_project_outputs" / "figures" / "10_top_centrality_comparison.png",
        "Merkezilik karsilastirma grafigi",
        6.2,
    )
    add_image(
        document,
        DOCS_DIR / "sna_project_outputs" / "figures" / "16_bridge_focus_compact.png",
        "Kopru dugumlere odaklanan alt ag",
        6.2,
    )

    frame_section(
        document,
        "Frame 5 - Topluluk ve Dayaniklilik",
        [
            "Yontem: Louvain community detection",
            "Topluluk sayisi: 8 | Modularity: 0.630465",
            "En buyuk topluluk: 69 dugum",
            "Topluluklar otomotiv, yemek_mutfak, egitim_kariyer ve anne_bebek gibi tema eksenlerinde dogal olarak ayrismistir.",
            "Tek bir yuksek degree dugumu cikinca ag dramatik sekilde parcalanmadi.",
            "Yuksek betweenness dugumleri hedefli cikarildiginda diameter 6'dan 7'ye cikti.",
        ],
        "Louvain sonucu, creatorlarin rastgele degil anlamli tema gruplari halinde toplandigini gosterdi. Dayaniklilik analizi ise bu topluluklari birbirine baglayan kopru dugumlerin kaybi halinde agin erisim verimliliginin bozuldugunu ortaya koydu.",
        DOCS_DIR / "sna_project_outputs" / "figures" / "07_louvain_communities_polished.png",
        "Topluluk yapisi",
        6.2,
    )
    add_image(
        document,
        DOCS_DIR / "sna_project_outputs" / "figures" / "14_robustness_comparison.png",
        "Dayaniklilik karsilastirmasi",
        6.2,
    )

    frame_section(
        document,
        "Frame 6 - Urun Odakli Onerme Mekanizmasi",
        [
            "1. Urun anlama katmani: keyword, kategori sinyali, kullanim baglami",
            "2. Graph tabanli secim katmani: relevance_score, authority_score, bridge_score, engagement_score, follower_score",
            "authority: tek kategoriye yakin urunler",
            "bridge: cok kategorili veya farkli topluluklari baglayan urunler",
        ],
        "Yani urun sinyali katmani talebi tanimlar, creator graph ise bu talebi yapisal olarak en uygun hesaplarla eslestirir. Projenin asil farki, sosyal ag analizini dogrudan onerme mekanizmasina baglamasidir.",
        DOCS_DIR / "excalidraw" / "exports" / "03_gercek_graph_modeli.png",
        "Urun sinyali ve graph baglantisi",
        6.3,
    )

    frame_section(
        document,
        "Frame 7 - Ornek Sorgular ve Sonuc",
        [
            "Ornek sorgular: kadin yuz serumu, spor matarasi, kamp termosu",
            "Net urunlerde authority modu one cikar.",
            "Yan kategori sinyali tasiyan urunlerde bridge / topluluk baglantilari deger kazanir.",
            "Sonuc: Proje sosyal ag analizi isterlerini karsilarken, bu analizleri urune gore mantikli influencer secen uygulamali bir sisteme donusturmektedir.",
        ],
        "Boylece proje sadece raporluk bir graph analizi degil; urun bazli karar uretebilen graph tabanli bir onerme sistemine donusmus oldu.",
        DOCS_DIR / "sna_project_outputs" / "figures" / "15_recommendation_examples.png",
        "Ornek urun sorgulari ve onerilen influencerlar",
        6.3,
    )

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_docx()
    print(output)
